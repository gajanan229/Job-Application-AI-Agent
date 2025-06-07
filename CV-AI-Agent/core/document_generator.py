#!/usr/bin/env python3
"""
Document Generation System for the Application Factory

This module handles DOCX template processing and PDF conversion.
Supports populating resume and cover letter templates with AI-generated content.
"""

import logging
import os
import tempfile
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pandas as pd

# Import for PDF conversion
try:
    from docx2pdf import convert
    PDF_CONVERSION_AVAILABLE = True
except ImportError:
    PDF_CONVERSION_AVAILABLE = False
    logging.warning("docx2pdf not available - PDF conversion will be disabled")

from config.settings import config
from config.logging_config import timing_decorator
from utils.error_handlers import DocumentGenerationError, FileProcessingError
from utils.file_utils import ensure_directory_exists
from utils.session_utils import SessionManager

# Configure logging
logger = logging.getLogger(__name__)

@dataclass
class DocumentData:
    """Data structure for holding document content"""
    content_type: str  # 'resume' or 'cover_letter'
    placeholders: Dict[str, str]  # mapping of placeholder -> content
    metadata: Dict[str, Any]
    
@dataclass
class GeneratedDocument:
    """Result of document generation"""
    content_type: str
    docx_path: str
    pdf_path: Optional[str]
    generation_time: float
    file_size_docx: int
    file_size_pdf: Optional[int]
    metadata: Dict[str, Any]

class DocumentTemplateManager:
    """Manages loading and processing of DOCX templates"""
    
    def __init__(self, template_dir: str = "Example CV"):
        """
        Initialize template manager
        
        Args:
            template_dir: Directory containing template files
        """
        self.template_dir = Path(template_dir)
        self.settings = config
        
        # Template file mappings
        self.template_files = {
            'resume': self.template_dir / 'resume_template.docx',
            'cover_letter': self.template_dir / 'cover_letter_template.docx'
        }
        
        # Validate templates exist
        self._validate_templates()
        
        # Template placeholders
        self.placeholders = {
            'resume': ['[SUMMARY]', '[PROJECT TITLE]', '[PROJECT BULLET POINTS]'],
            'cover_letter': ['[INTRODUCTION]', '[BODY]', '[CONCLUSION]']
        }
    
    def _validate_templates(self) -> None:
        """Validate that all required templates exist"""
        for doc_type, template_path in self.template_files.items():
            if not template_path.exists():
                raise FileProcessingError(f"Template file not found: {template_path}")
            
            try:
                # Test loading the template
                Document(str(template_path))
                logger.info(f"Validated template: {template_path}")
            except Exception as e:
                raise DocumentGenerationError(f"Invalid template file {template_path}: {e}")
    
    @timing_decorator
    def load_template(self, doc_type: str) -> Document:
        """
        Load a DOCX template
        
        Args:
            doc_type: Type of document ('resume' or 'cover_letter')
            
        Returns:
            Loaded Document object
            
        Raises:
            DocumentGenerationError: If template cannot be loaded
        """
        if doc_type not in self.template_files:
            raise DocumentGenerationError(f"Unknown document type: {doc_type}")
        
        template_path = self.template_files[doc_type]
        
        try:
            doc = Document(str(template_path))
            logger.info(f"Loaded template: {template_path}")
            return doc
        except Exception as e:
            raise DocumentGenerationError(f"Failed to load template {template_path}: {e}")
    
    def get_placeholders(self, doc_type: str) -> List[str]:
        """Get list of placeholders for a document type"""
        return self.placeholders.get(doc_type, [])

class DocumentPopulator:
    """Handles populating templates with content"""
    
    def __init__(self):
        self.settings = config
    
    @timing_decorator
    def populate_resume(self, doc: Document, content: Dict[str, str]) -> Document:
        """
        Populate resume template with generated content
        
        Args:
            doc: Document object to populate
            content: Dictionary with keys 'summary', 'project_title', 'project_bullets'
            
        Returns:
            Populated document
        """
        try:
            # Mapping of placeholders to content keys
            placeholder_mapping = {
                '[SUMMARY]': content.get('summary', ''),
                '[PROJECT TITLE]': content.get('project_title', ''),
                '[PROJECT BULLET POINTS]': content.get('project_bullets', '')
            }
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, replacement in placeholder_mapping.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)
                        logger.debug(f"Replaced {placeholder} in paragraph")
            
            # Also check in tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, replacement in placeholder_mapping.items():
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, replacement)
                                    logger.debug(f"Replaced {placeholder} in table cell")
            
            logger.info("Successfully populated resume template")
            return doc
            
        except Exception as e:
            raise DocumentGenerationError(f"Failed to populate resume: {e}")
    
    @timing_decorator
    def populate_cover_letter(self, doc: Document, content: Dict[str, str]) -> Document:
        """
        Populate cover letter template with generated content
        
        Args:
            doc: Document object to populate
            content: Dictionary with keys 'introduction', 'body', 'conclusion'
            
        Returns:
            Populated document
        """
        try:
            # Mapping of placeholders to content keys
            placeholder_mapping = {
                '[INTRODUCTION]': content.get('introduction', ''),
                '[BODY]': content.get('body', ''),
                '[CONCLUSION]': content.get('conclusion', '')
            }
            
            # Replace placeholders in paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, replacement in placeholder_mapping.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, replacement)
                        logger.debug(f"Replaced {placeholder} in paragraph")
            
            # Also check in tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for placeholder, replacement in placeholder_mapping.items():
                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, replacement)
                                    logger.debug(f"Replaced {placeholder} in table cell")
            
            logger.info("Successfully populated cover letter template")
            return doc
            
        except Exception as e:
            raise DocumentGenerationError(f"Failed to populate cover letter: {e}")

class PDFConverter:
    """Handles DOCX to PDF conversion"""
    
    def __init__(self):
        self.conversion_available = PDF_CONVERSION_AVAILABLE
        
    @timing_decorator
    def convert_to_pdf(self, docx_path: str, pdf_path: Optional[str] = None) -> Optional[str]:
        """
        Convert DOCX file to PDF
        
        Args:
            docx_path: Path to DOCX file
            pdf_path: Optional output PDF path (if None, uses same name with .pdf extension)
            
        Returns:
            Path to generated PDF file, or None if conversion failed
            
        Raises:
            DocumentGenerationError: If conversion fails
        """
        if not self.conversion_available:
            logger.warning("PDF conversion not available - docx2pdf not installed")
            return None
        
        try:
            # Determine output path
            if pdf_path is None:
                pdf_path = Path(docx_path).with_suffix('.pdf')
            
            # Ensure output directory exists
            ensure_directory_exists(Path(pdf_path).parent)
            
            # Convert using docx2pdf
            convert(docx_path, str(pdf_path))
            
            if Path(pdf_path).exists():
                logger.info(f"Successfully converted {docx_path} to {pdf_path}")
                return str(pdf_path)
            else:
                raise DocumentGenerationError("PDF file was not created")
                
        except Exception as e:
            logger.error(f"PDF conversion failed: {e}")
            raise DocumentGenerationError(f"Failed to convert to PDF: {e}")

class DocumentGenerator:
    """Main document generation orchestrator"""
    
    def __init__(self, output_dir: str = "generated_documents"):
        """
        Initialize document generator
        
        Args:
            output_dir: Directory for saving generated documents
        """
        self.output_dir = Path(output_dir)
        self.template_manager = DocumentTemplateManager()
        self.populator = DocumentPopulator()
        self.pdf_converter = PDFConverter()
        self.settings = config
        
        # Ensure output directory exists
        ensure_directory_exists(self.output_dir)
    
    def _generate_filename(self, doc_type: str, file_format: str) -> str:
        """Generate unique filename for output documents"""
        timestamp = SessionManager.get_timestamp()
        return f"{doc_type}_{timestamp}.{file_format}"
    
    @timing_decorator
    def generate_resume(self, content: Dict[str, str], 
                       output_filename: Optional[str] = None) -> GeneratedDocument:
        """
        Generate a complete resume document
        
        Args:
            content: Dictionary with 'summary', 'project_title', 'project_bullets'
            output_filename: Optional custom filename (without extension)
            
        Returns:
            GeneratedDocument with paths and metadata
        """
        start_time = datetime.now()
        
        try:
            # Load template
            doc = self.template_manager.load_template('resume')
            
            # Populate with content
            doc = self.populator.populate_resume(doc, content)
            
            # Generate filenames
            if output_filename:
                docx_filename = f"{output_filename}.docx"
                pdf_filename = f"{output_filename}.pdf"
            else:
                docx_filename = self._generate_filename('resume', 'docx')
                pdf_filename = self._generate_filename('resume', 'pdf')
            
            docx_path = self.output_dir / docx_filename
            pdf_path = self.output_dir / pdf_filename
            
            # Save DOCX
            doc.save(str(docx_path))
            docx_size = docx_path.stat().st_size
            
            # Convert to PDF
            pdf_output_path = self.pdf_converter.convert_to_pdf(str(docx_path), str(pdf_path))
            pdf_size = None
            if pdf_output_path and Path(pdf_output_path).exists():
                pdf_size = Path(pdf_output_path).stat().st_size
            
            generation_time = (datetime.now() - start_time).total_seconds()
            
            result = GeneratedDocument(
                content_type='resume',
                docx_path=str(docx_path),
                pdf_path=pdf_output_path,
                generation_time=generation_time,
                file_size_docx=docx_size,
                file_size_pdf=pdf_size,
                metadata={
                    'template_used': 'resume_template.docx',
                    'placeholders_filled': list(content.keys()),
                    'generated_at': datetime.now().isoformat(),
                    'pdf_conversion_available': self.pdf_converter.conversion_available
                }
            )
            
            logger.info(f"Generated resume: {docx_path} (PDF: {pdf_output_path})")
            return result
            
        except Exception as e:
            raise DocumentGenerationError(f"Failed to generate resume: {e}")
    
    @timing_decorator
    def generate_cover_letter(self, content: Dict[str, str], 
                             output_filename: Optional[str] = None) -> GeneratedDocument:
        """
        Generate a complete cover letter document
        
        Args:
            content: Dictionary with keys 'introduction', 'body', 'conclusion'
            output_filename: Optional custom filename (without extension)
            
        Returns:
            GeneratedDocument with paths and metadata
        """
        start_time = datetime.now()
        
        try:
            # Load template
            doc = self.template_manager.load_template('cover_letter')
            
            # Populate with content
            doc = self.populator.populate_cover_letter(doc, content)
            
            # Generate filenames
            if output_filename:
                docx_filename = f"{output_filename}.docx"
                pdf_filename = f"{output_filename}.pdf"
            else:
                docx_filename = self._generate_filename('cover_letter', 'docx')
                pdf_filename = self._generate_filename('cover_letter', 'pdf')
            
            docx_path = self.output_dir / docx_filename
            pdf_path = self.output_dir / pdf_filename
            
            # Save DOCX
            doc.save(str(docx_path))
            docx_size = docx_path.stat().st_size
            
            # Convert to PDF
            pdf_output_path = self.pdf_converter.convert_to_pdf(str(docx_path), str(pdf_path))
            pdf_size = None
            if pdf_output_path and Path(pdf_output_path).exists():
                pdf_size = Path(pdf_output_path).stat().st_size
            
            generation_time = (datetime.now() - start_time).total_seconds()
            
            result = GeneratedDocument(
                content_type='cover_letter',
                docx_path=str(docx_path),
                pdf_path=pdf_output_path,
                generation_time=generation_time,
                file_size_docx=docx_size,
                file_size_pdf=pdf_size,
                metadata={
                    'template_used': 'cover_letter_template.docx',
                    'placeholders_filled': list(content.keys()),
                    'generated_at': datetime.now().isoformat(),
                    'pdf_conversion_available': self.pdf_converter.conversion_available
                }
            )
            
            logger.info(f"Generated cover letter: {docx_path} (PDF: {pdf_output_path})")
            return result
            
        except Exception as e:
            raise DocumentGenerationError(f"Failed to generate cover letter: {e}")
    
    def generate_both(self, resume_content: Dict[str, str], 
                      cover_letter_content: Dict[str, str],
                      output_prefix: Optional[str] = None) -> Tuple[GeneratedDocument, GeneratedDocument]:
        """
        Generate both resume and cover letter
        
        Args:
            resume_content: Content for resume
            cover_letter_content: Content for cover letter
            output_prefix: Optional prefix for output files
            
        Returns:
            Tuple of (resume_document, cover_letter_document)
        """
        try:
            # Generate both documents
            if output_prefix:
                resume_filename = f"{output_prefix}_resume"
                cover_letter_filename = f"{output_prefix}_cover_letter"
            else:
                resume_filename = None
                cover_letter_filename = None
            
            resume_doc = self.generate_resume(resume_content, resume_filename)
            cover_letter_doc = self.generate_cover_letter(cover_letter_content, cover_letter_filename)
            
            logger.info("Successfully generated both resume and cover letter")
            return resume_doc, cover_letter_doc
            
        except Exception as e:
            raise DocumentGenerationError(f"Failed to generate documents: {e}")
    
    def get_template_info(self) -> Dict[str, Any]:
        """Get information about available templates"""
        return {
            'templates': {
                doc_type: {
                    'path': str(self.template_manager.template_files[doc_type]),
                    'placeholders': self.template_manager.get_placeholders(doc_type),
                    'exists': self.template_manager.template_files[doc_type].exists()
                }
                for doc_type in self.template_manager.template_files.keys()
            },
            'output_directory': str(self.output_dir),
            'pdf_conversion_available': self.pdf_converter.conversion_available
        }

# Convenience functions for direct use
def load_docx_template(template_path: str) -> Document:
    """Load a DOCX template file"""
    template_manager = DocumentTemplateManager()
    return template_manager.load_template(Path(template_path).stem.replace('_template', ''))

def populate_resume_docx(doc: Document, content: Dict[str, str]) -> Document:
    """Populate resume template with content"""
    populator = DocumentPopulator()
    return populator.populate_resume(doc, content)

def populate_cover_letter_docx(doc: Document, content: Dict[str, str]) -> Document:
    """Populate cover letter template with content"""
    populator = DocumentPopulator()
    return populator.populate_cover_letter(doc, content)

def convert_docx_to_pdf(docx_path: str, pdf_path: Optional[str] = None) -> Optional[str]:
    """Convert DOCX to PDF"""
    converter = PDFConverter()
    return converter.convert_to_pdf(docx_path, pdf_path)

# Export main classes and functions
__all__ = [
    'DocumentGenerator',
    'DocumentTemplateManager', 
    'DocumentPopulator',
    'PDFConverter',
    'DocumentData',
    'GeneratedDocument',
    'load_docx_template',
    'populate_resume_docx',
    'populate_cover_letter_docx',
    'convert_docx_to_pdf'
] 