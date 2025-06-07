#!/usr/bin/env python3
"""
Tests for the Document Generation System
"""

import pytest
import tempfile
import shutil
from pathlib import Path
from unittest.mock import Mock, patch, MagicMock
from docx import Document

from core.document_generator import (
    DocumentGenerator,
    DocumentTemplateManager,
    DocumentPopulator,
    PDFConverter,
    DocumentData,
    GeneratedDocument,
    load_docx_template,
    populate_resume_docx,
    populate_cover_letter_docx,
    convert_docx_to_pdf
)
from utils.error_handlers import DocumentGenerationError, FileProcessingError


class TestDocumentTemplateManager:
    """Test the DocumentTemplateManager class"""
    
    def test_init_with_valid_templates(self):
        """Test initialization with valid template directory"""
        manager = DocumentTemplateManager("Example CV")
        
        assert manager.template_dir == Path("Example CV")
        assert 'resume' in manager.template_files
        assert 'cover_letter' in manager.template_files
        assert len(manager.placeholders['resume']) == 3
        assert len(manager.placeholders['cover_letter']) == 3
    
    def test_init_with_missing_templates(self):
        """Test initialization with missing template directory"""
        with pytest.raises(FileProcessingError):
            DocumentTemplateManager("nonexistent_dir")
    
    def test_load_template_resume(self):
        """Test loading resume template"""
        manager = DocumentTemplateManager("Example CV")
        doc = manager.load_template('resume')
        
        assert doc is not None
        assert hasattr(doc, 'paragraphs')  # Check it's a Document-like object
    
    def test_load_template_cover_letter(self):
        """Test loading cover letter template"""
        manager = DocumentTemplateManager("Example CV")
        doc = manager.load_template('cover_letter')
        
        assert doc is not None
        assert hasattr(doc, 'paragraphs')  # Check it's a Document-like object
    
    def test_load_template_invalid_type(self):
        """Test loading invalid template type"""
        manager = DocumentTemplateManager("Example CV")
        
        with pytest.raises(DocumentGenerationError):
            manager.load_template('invalid_type')
    
    def test_get_placeholders(self):
        """Test getting placeholders for document types"""
        manager = DocumentTemplateManager("Example CV")
        
        resume_placeholders = manager.get_placeholders('resume')
        cover_letter_placeholders = manager.get_placeholders('cover_letter')
        
        assert '[SUMMARY]' in resume_placeholders
        assert '[PROJECT TITLE]' in resume_placeholders
        assert '[PROJECT BULLET POINTS]' in resume_placeholders
        
        assert '[INTRODUCTION]' in cover_letter_placeholders
        assert '[BODY]' in cover_letter_placeholders
        assert '[CONCLUSION]' in cover_letter_placeholders


class TestDocumentPopulator:
    """Test the DocumentPopulator class"""
    
    def setup_method(self):
        """Set up test fixtures"""
        self.populator = DocumentPopulator()
        self.template_manager = DocumentTemplateManager("Example CV")
    
    def test_populate_resume_basic(self):
        """Test basic resume population"""
        doc = self.template_manager.load_template('resume')
        content = {
            'summary': 'Test summary content',
            'project_title': 'Test Project',
            'project_bullets': '• Test bullet point 1\n• Test bullet point 2'
        }
        
        populated_doc = self.populator.populate_resume(doc, content)
        
        # Check that placeholders were replaced
        doc_text = '\n'.join([p.text for p in populated_doc.paragraphs])
        assert 'Test summary content' in doc_text
        assert 'Test Project' in doc_text
        assert 'Test bullet point 1' in doc_text
        assert '[SUMMARY]' not in doc_text
        assert '[PROJECT TITLE]' not in doc_text
        assert '[PROJECT BULLET POINTS]' not in doc_text
    
    def test_populate_cover_letter_basic(self):
        """Test basic cover letter population"""
        doc = self.template_manager.load_template('cover_letter')
        content = {
            'introduction': 'Dear Hiring Manager,',
            'body': 'This is the body of the cover letter.',
            'conclusion': 'Thank you for your consideration.'
        }
        
        populated_doc = self.populator.populate_cover_letter(doc, content)
        
        # Check that placeholders were replaced
        doc_text = '\n'.join([p.text for p in populated_doc.paragraphs])
        assert 'Dear Hiring Manager,' in doc_text
        assert 'This is the body of the cover letter.' in doc_text
        assert 'Thank you for your consideration.' in doc_text
        assert '[INTRODUCTION]' not in doc_text
        assert '[BODY]' not in doc_text
        assert '[CONCLUSION]' not in doc_text
    
    def test_populate_resume_missing_content(self):
        """Test resume population with missing content keys"""
        doc = self.template_manager.load_template('resume')
        content = {'summary': 'Only summary provided'}
        
        populated_doc = self.populator.populate_resume(doc, content)
        
        # Check that available content was used, missing keys were replaced with empty strings
        doc_text = '\n'.join([p.text for p in populated_doc.paragraphs])
        assert 'Only summary provided' in doc_text
        assert '[SUMMARY]' not in doc_text
        # Other placeholders should be replaced with empty strings
        assert '[PROJECT TITLE]' not in doc_text
        assert '[PROJECT BULLET POINTS]' not in doc_text
    
    def test_populate_cover_letter_missing_content(self):
        """Test cover letter population with missing content keys"""
        doc = self.template_manager.load_template('cover_letter')
        content = {'introduction': 'Only introduction provided'}
        
        populated_doc = self.populator.populate_cover_letter(doc, content)
        
        # Check that available content was used, missing keys were replaced with empty strings
        doc_text = '\n'.join([p.text for p in populated_doc.paragraphs])
        assert 'Only introduction provided' in doc_text
        assert '[INTRODUCTION]' not in doc_text
        # Other placeholders should be replaced with empty strings
        assert '[BODY]' not in doc_text
        assert '[CONCLUSION]' not in doc_text


class TestPDFConverter:
    """Test the PDFConverter class"""
    
    def setup_method(self):
        """Set up test fixtures"""
        self.converter = PDFConverter()
        self.temp_dir = tempfile.mkdtemp()
        self.temp_docx = Path(self.temp_dir) / "test.docx"
        self.temp_pdf = Path(self.temp_dir) / "test.pdf"
        
        # Create a simple test DOCX file
        doc = Document()
        doc.add_paragraph("Test document content")
        doc.save(str(self.temp_docx))
    
    def teardown_method(self):
        """Clean up test fixtures"""
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    @patch('core.document_generator.PDF_CONVERSION_AVAILABLE', True)
    @patch('core.document_generator.convert')
    def test_convert_to_pdf_success(self, mock_convert):
        """Test successful PDF conversion"""
        # Mock the conversion to create the PDF file
        def create_pdf(docx_path, pdf_path):
            Path(pdf_path).touch()
        
        mock_convert.side_effect = create_pdf
        
        result = self.converter.convert_to_pdf(str(self.temp_docx), str(self.temp_pdf))
        
        assert result == str(self.temp_pdf)
        mock_convert.assert_called_once()
    
    @patch('core.document_generator.PDF_CONVERSION_AVAILABLE', False)
    def test_convert_to_pdf_unavailable(self):
        """Test PDF conversion when not available"""
        converter = PDFConverter()
        result = converter.convert_to_pdf(str(self.temp_docx))
        
        assert result is None
    
    @patch('core.document_generator.PDF_CONVERSION_AVAILABLE', True)
    @patch('core.document_generator.convert')
    def test_convert_to_pdf_failure(self, mock_convert):
        """Test PDF conversion failure"""
        mock_convert.side_effect = Exception("Conversion failed")
        
        with pytest.raises(DocumentGenerationError):
            self.converter.convert_to_pdf(str(self.temp_docx))
    
    def test_convert_to_pdf_auto_path(self):
        """Test PDF conversion with automatic path generation"""
        # This test will skip if PDF conversion is not available
        if not self.converter.conversion_available:
            pytest.skip("PDF conversion not available")
        
        with patch('core.document_generator.convert') as mock_convert:
            def create_pdf(docx_path, pdf_path):
                Path(pdf_path).touch()
            
            mock_convert.side_effect = create_pdf
            
            result = self.converter.convert_to_pdf(str(self.temp_docx))
            
            expected_pdf = str(self.temp_docx.with_suffix('.pdf'))
            assert result == expected_pdf


class TestDocumentGenerator:
    """Test the DocumentGenerator class"""
    
    def setup_method(self):
        """Set up test fixtures"""
        self.temp_dir = tempfile.mkdtemp()
        self.generator = DocumentGenerator(output_dir=self.temp_dir)
    
    def teardown_method(self):
        """Clean up test fixtures"""
        shutil.rmtree(self.temp_dir, ignore_errors=True)
    
    def test_init(self):
        """Test DocumentGenerator initialization"""
        assert self.generator.output_dir == Path(self.temp_dir)
        assert isinstance(self.generator.template_manager, DocumentTemplateManager)
        assert isinstance(self.generator.populator, DocumentPopulator)
        assert isinstance(self.generator.pdf_converter, PDFConverter)
    
    def test_generate_filename(self):
        """Test filename generation"""
        filename = self.generator._generate_filename('resume', 'docx')
        
        assert filename.startswith('resume_')
        assert filename.endswith('.docx')
        assert len(filename) > 10  # Should include timestamp
    
    @patch('core.document_generator.PDFConverter.convert_to_pdf')
    def test_generate_resume(self, mock_pdf_convert):
        """Test resume generation"""
        mock_pdf_convert.return_value = str(Path(self.temp_dir) / "test_resume.pdf")
        
        content = {
            'summary': 'Test summary',
            'project_title': 'Test Project',
            'project_bullets': '• Test bullet'
        }
        
        result = self.generator.generate_resume(content, 'test_resume')
        
        assert isinstance(result, GeneratedDocument)
        assert result.content_type == 'resume'
        assert Path(result.docx_path).exists()
        assert result.pdf_path is not None
        assert result.generation_time > 0
        assert result.file_size_docx > 0
    
    @patch('core.document_generator.PDFConverter.convert_to_pdf')
    def test_generate_cover_letter(self, mock_pdf_convert):
        """Test cover letter generation"""
        mock_pdf_convert.return_value = str(Path(self.temp_dir) / "test_cover_letter.pdf")
        
        content = {
            'introduction': 'Dear Hiring Manager,',
            'body': 'Cover letter body',
            'conclusion': 'Thank you'
        }
        
        result = self.generator.generate_cover_letter(content, 'test_cover_letter')
        
        assert isinstance(result, GeneratedDocument)
        assert result.content_type == 'cover_letter'
        assert Path(result.docx_path).exists()
        assert result.pdf_path is not None
        assert result.generation_time > 0
        assert result.file_size_docx > 0
    
    @patch('core.document_generator.PDFConverter.convert_to_pdf')
    def test_generate_both(self, mock_pdf_convert):
        """Test generating both documents"""
        mock_pdf_convert.return_value = "dummy_path.pdf"
        
        resume_content = {
            'summary': 'Test summary',
            'project_title': 'Test Project',
            'project_bullets': '• Test bullet'
        }
        
        cover_letter_content = {
            'introduction': 'Dear Hiring Manager,',
            'body': 'Cover letter body',
            'conclusion': 'Thank you'
        }
        
        resume_doc, cover_letter_doc = self.generator.generate_both(
            resume_content, cover_letter_content, 'test_batch'
        )
        
        assert isinstance(resume_doc, GeneratedDocument)
        assert isinstance(cover_letter_doc, GeneratedDocument)
        assert resume_doc.content_type == 'resume'
        assert cover_letter_doc.content_type == 'cover_letter'
        assert Path(resume_doc.docx_path).exists()
        assert Path(cover_letter_doc.docx_path).exists()
    
    def test_get_template_info(self):
        """Test getting template information"""
        info = self.generator.get_template_info()
        
        assert 'templates' in info
        assert 'output_directory' in info
        assert 'pdf_conversion_available' in info
        assert 'resume' in info['templates']
        assert 'cover_letter' in info['templates']


class TestConvenienceFunctions:
    """Test the convenience functions"""
    
    def test_load_docx_template(self):
        """Test the load_docx_template convenience function"""
        doc = load_docx_template("Example CV/resume_template.docx")
        assert doc is not None
        assert hasattr(doc, 'paragraphs')  # Check it's a Document-like object
    
    def test_populate_resume_docx(self):
        """Test the populate_resume_docx convenience function"""
        doc = load_docx_template("Example CV/resume_template.docx")
        content = {
            'summary': 'Test summary',
            'project_title': 'Test Project',
            'project_bullets': '• Test bullet'
        }
        
        populated_doc = populate_resume_docx(doc, content)
        assert populated_doc is not None
        assert hasattr(populated_doc, 'paragraphs')  # Check it's a Document-like object
        
        doc_text = '\n'.join([p.text for p in populated_doc.paragraphs])
        assert 'Test summary' in doc_text
    
    def test_populate_cover_letter_docx(self):
        """Test the populate_cover_letter_docx convenience function"""
        doc = load_docx_template("Example CV/cover_letter_template.docx")
        content = {
            'introduction': 'Dear Hiring Manager,',
            'body': 'Cover letter body',
            'conclusion': 'Thank you'
        }
        
        populated_doc = populate_cover_letter_docx(doc, content)
        assert populated_doc is not None
        assert hasattr(populated_doc, 'paragraphs')  # Check it's a Document-like object
        
        doc_text = '\n'.join([p.text for p in populated_doc.paragraphs])
        assert 'Dear Hiring Manager,' in doc_text
    
    @patch('core.document_generator.PDFConverter.convert_to_pdf')
    def test_convert_docx_to_pdf(self, mock_convert):
        """Test the convert_docx_to_pdf convenience function"""
        mock_convert.return_value = "test.pdf"
        
        result = convert_docx_to_pdf("test.docx")
        assert result == "test.pdf"
        mock_convert.assert_called_once()


class TestDocumentData:
    """Test the DocumentData dataclass"""
    
    def test_document_data_creation(self):
        """Test creating DocumentData instance"""
        data = DocumentData(
            content_type="resume",
            placeholders={"[SUMMARY]": "Test summary"},
            metadata={"test": "value"}
        )
        
        assert data.content_type == "resume"
        assert data.placeholders["[SUMMARY]"] == "Test summary"
        assert data.metadata["test"] == "value"


class TestGeneratedDocument:
    """Test the GeneratedDocument dataclass"""
    
    def test_generated_document_creation(self):
        """Test creating GeneratedDocument instance"""
        doc = GeneratedDocument(
            content_type="resume",
            docx_path="/path/to/resume.docx",
            pdf_path="/path/to/resume.pdf",
            generation_time=2.5,
            file_size_docx=50000,
            file_size_pdf=45000,
            metadata={"test": "value"}
        )
        
        assert doc.content_type == "resume"
        assert doc.docx_path == "/path/to/resume.docx"
        assert doc.pdf_path == "/path/to/resume.pdf"
        assert doc.generation_time == 2.5
        assert doc.file_size_docx == 50000
        assert doc.file_size_pdf == 45000
        assert doc.metadata["test"] == "value"


if __name__ == "__main__":
    pytest.main([__file__]) 